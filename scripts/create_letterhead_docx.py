#!/usr/bin/env python3
"""Generate Pneuma Health letterhead .docx (logo + brand colors from site palette)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE = Path(__file__).resolve().parent.parent
LOGO = BASE / "assets" / "logo-pneuma.png"
OUT = BASE / "Pneuma-Health-Letterhead.docx"

# OOXML fill without # — matches styles.css :root
NAVY = "426B96"
SAGE = "819F96"
PERI = "727FB7"


def set_paragraph_shading(paragraph, color_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    p_pr.append(shd)


def clear_header(header) -> None:
    el = header._element
    for p in el.findall(qn("w:p")):
        el.remove(p)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    header = section.header
    clear_header(header)

    bar = header.add_paragraph()
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after = Pt(0)
    bar.paragraph_format.line_spacing = Pt(11)
    set_paragraph_shading(bar, NAVY)
    bar.add_run("\u00a0").font.size = Pt(4)

    p_logo = header.add_paragraph()
    p_logo.paragraph_format.space_before = Pt(10)
    p_logo.paragraph_format.space_after = Pt(2)
    if LOGO.is_file():
        p_logo.add_run().add_picture(str(LOGO), width=Inches(2.35))
    else:
        r = p_logo.add_run("Pneuma Health")
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = RGBColor(0x42, 0x6B, 0x96)

    p_name = header.add_paragraph()
    r = p_name.add_run("Pneuma Health Inc.")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x42, 0x6B, 0x96)

    p_tag = header.add_paragraph()
    r = p_tag.add_run("Integrative mental health care")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x81, 0x9F, 0x96)

    line = header.add_paragraph()
    line.paragraph_format.space_before = Pt(6)
    line.paragraph_format.space_after = Pt(0)
    line.paragraph_format.line_spacing = Pt(4)
    set_paragraph_shading(line, PERI)
    line.add_run("\u00a0").font.size = Pt(2)

    # Footer
    footer = section.footer
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)
    def footer_line(text: str, *, italic: bool = False) -> None:
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(text)
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor(0x42, 0x6B, 0x96) if not italic else RGBColor(
            0x81, 0x9F, 0x96
        )
        rr.italic = italic

    footer_line("info@pneumahealth.ca  ·  pneumahealth.ca")
    footer_line(
        "3230 Yonge Street #4092, Toronto, ON M4N 3P6, Canada",
        italic=True,
    )
    footer_line(
        "Hours: Mon–Fri 10am–4pm  ·  Sat 10am–2pm",
        italic=True,
    )
    footer.paragraphs[0].paragraph_format.space_before = Pt(6)

    # Body starter (user deletes or keeps)
    doc.add_paragraph()
    doc.add_paragraph()
    p_date = doc.add_paragraph()
    p_date.add_run("March 28, 2026").font.size = Pt(11)

    doc.add_paragraph()
    p_dear = doc.add_paragraph()
    p_dear.add_run("Dear ").font.size = Pt(11)
    p_dear.add_run("[Name]").font.size = Pt(11)

    doc.add_paragraph()
    body = doc.add_paragraph()
    body.add_run(
        "Your letter begins here. This template uses your site colors: navy (#426b96), "
        "sage (#819f96), and periwinkle (#727fb7) in the header and footer."
    ).font.size = Pt(11)

    doc.add_paragraph()
    close = doc.add_paragraph()
    close.add_run("Sincerely,").font.size = Pt(11)
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("[Your name]").font.size = Pt(11)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
